from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE = Path(r"C:\Users\chiqu\Downloads\DocumentacionFinal (1) (1).docx")
OUTPUT = Path(r"C:\Users\chiqu\tecwebSistemaGym\SistemaGym.Api\DocumentacionFinal_SistemaGym_Actualizada.docx")

BASE_URL = "https://sistemagym-api-gee7bwdvepgzc0f6.canadacentral-01.azurewebsites.net"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "8EA9C1")


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
            if row_index == 0:
                set_cell_shading(cell, "D9EAF7")


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    new_para._element = new_p
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(paragraph, rows, cols):
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Inches(6.3))
    paragraph._p.addnext(table._tbl)
    return table


doc = Document(SOURCE)

for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    if text == "Abril del 2026":
        paragraph.text = "Mayo del 2026"
    elif text == "Las contraseñas de usuarios deben implementar hash (bcrypt) en versiones futuras.":
        paragraph.text = "Las contraseñas de usuarios se almacenan con hashing PBKDF2, salt aleatorio y comparación segura mediante PasswordService."
    elif text == "Se recomienda implementar autenticación JWT para versiones futuras del sistema.":
        paragraph.text = "La API implementa autenticación JWT Bearer, validando issuer, audience, lifetime y clave de firma. Los endpoints principales están protegidos con [Authorize]."
    elif text == "Base URL local: http://localhost:5227":
        paragraph.text = f"Base URL pública en Azure: {BASE_URL}"
    elif text.startswith("Swagger / OpenAPI: http://localhost:5227/swagger"):
        paragraph.text = f"Swagger / OpenAPI publicado: {BASE_URL}/swagger"
    elif text == "Requests realizados en Postman para los casos de uso CU-05 a CU-08:":
        paragraph.text = "Requests realizados en Swagger/Postman para los casos de uso CU-01 a CU-10:"
    elif text == "POST http://localhost:5227/api/PlanMembresia/dto/mapper":
        paragraph.text = f"POST {BASE_URL}/api/PlanMembresia/dto/mapper"


anchor = None
for paragraph in doc.paragraphs:
    if paragraph.text.strip() == "6. Anexos":
        anchor = paragraph
        break

if anchor is not None:
    current = insert_paragraph_after(anchor, "", None)

    current = insert_paragraph_after(current, "5.1 Estado Actual de Criterios de Evaluación", "Heading 2")
    current = insert_paragraph_after(
        current,
        "El proyecto fue revisado contra los criterios actuales de evaluación. La API compila correctamente y cuenta con publicación en Azure App Service, documentación Swagger, autenticación JWT, registro/login de usuarios, hashing de contraseñas, mensajes estructurados y paginación en consultas principales."
    )

    criteria = [
        ("Criterio", "Estado", "Evidencia en el proyecto"),
        ("Mensajes", "Cumple", "ApiResponse<T>, Message, ResponseData y respuestas desde servicios GetAll...ResponseAsync."),
        ("Paginación", "Cumple", "PagedList, Pagination, PaginationQueryFilter y parámetros pageNumber/pageSize."),
        ("Swagger", "Cumple", "AddSwaggerGen, UseSwagger y UseSwaggerUI publicados en /swagger."),
        ("Seguridad API con JWT", "Cumple", "AddAuthentication().AddJwtBearer(), [Authorize] y validación de issuer/audience/secret."),
        ("Registrar usuario", "Cumple", "POST /api/Security registra usuarios con rol y contraseña hasheada."),
        ("Login user", "Cumple", "POST /api/Token valida credenciales y genera JWT."),
        ("Generate hashing passwords", "Cumple", "PasswordService usa PBKDF2 con SHA256, salt aleatorio e iteraciones configurables."),
        ("Publicado en Azure", "Cumple", f"API disponible en {BASE_URL} y Swagger en {BASE_URL}/swagger."),
    ]

    table = insert_table_after(current, len(criteria), 3)
    for row_index, row in enumerate(criteria):
        for col_index, value in enumerate(row):
            set_cell_text(table.cell(row_index, col_index), value, bold=(row_index == 0))
    style_table(table)
    current = insert_paragraph_after(current, "")

    current = insert_paragraph_after(current, "5.2 Endpoints Recomendados para los Casos de Uso", "Heading 2")
    endpoint_rows = [
        ("Caso de uso", "Endpoint recomendado", "Método"),
        ("CU-01 Registrar Cliente", "/api/Cliente/dto/mapper", "POST"),
        ("CU-02 Consultar Clientes", "/api/Cliente/dto/mapper?pageNumber=1&pageSize=10", "GET"),
        ("CU-03 Actualizar Cliente", "/api/Cliente/dto/mapper/{id}", "PUT"),
        ("CU-04 Eliminar Cliente", "/api/Cliente/dto/{id}", "DELETE"),
        ("CU-05 Crear Plan de Membresía", "/api/PlanMembresia/dto/mapper", "POST"),
        ("CU-06 Consultar Planes de Membresía", "/api/PlanMembresia/dto/mapper?pageNumber=1&pageSize=10", "GET"),
        ("CU-07 Asignar Membresía a Cliente", "/api/Membresia/dto/mapper", "POST"),
        ("CU-08 Consultar Membresías", "/api/Membresia/dto/mapper?pageNumber=1&pageSize=10", "GET"),
        ("CU-09 Registrar Pago de Membresía", "/api/Pago/dto/mapper", "POST"),
        ("CU-10 Consultar Pagos de Membresía", "/api/Pago/dto/mapper?pageNumber=1&pageSize=10", "GET"),
    ]
    endpoint_table = insert_table_after(current, len(endpoint_rows), 3)
    for row_index, row in enumerate(endpoint_rows):
        for col_index, value in enumerate(row):
            set_cell_text(endpoint_table.cell(row_index, col_index), value, bold=(row_index == 0))
    style_table(endpoint_table)
    current = insert_paragraph_after(current, "")

    current = insert_paragraph_after(current, "5.3 Flujo de Seguridad para Pruebas", "Heading 2")
    security_steps = [
        "Abrir Swagger en la URL pública /swagger.",
        "Ejecutar POST /api/Token con user=admin@gym.com y password=123456.",
        "Copiar el token devuelto por la API.",
        "Presionar Authorize en Swagger y registrar el valor con el formato Bearer {token}.",
        "Ejecutar los endpoints protegidos de clientes, planes, membresías y pagos.",
    ]
    for index, step in list(enumerate(security_steps, start=1))[::-1]:
        p = insert_paragraph_after(current, f"{index}. {step}")
        current = p

    note = insert_paragraph_after(
        current,
        "Nota: la contraseña del usuario administrador se guarda en la tabla usuario como hash generado por PasswordService, no como texto plano."
    )
    note.runs[0].italic = True

    # Move the inserted block before Anexos by reordering XML nodes.
    inserted_nodes = []
    node = anchor._p.getnext()
    while node is not None:
        inserted_nodes.append(node)
        node = node.getnext()
    for node in inserted_nodes:
        anchor._p.addprevious(deepcopy(node))
    for node in inserted_nodes:
        node.getparent().remove(node)


for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.text = "Sistema Gym API - Documentacion actualizada para despliegue en Azure"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(OUTPUT)
print(OUTPUT)
