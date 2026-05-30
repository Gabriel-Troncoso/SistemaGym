from pathlib import Path
from docx import Document


source = Path(r"C:\Users\chiqu\Downloads\DocumentacionFinal (1) (1).docx")
doc = Document(source)

print("PARAGRAPHS")
for i, paragraph in enumerate(doc.paragraphs, start=1):
    text = paragraph.text.strip()
    if text:
        print(f"{i:04d} [{paragraph.style.name}] {text}")

print("\nTABLES")
for table_index, table in enumerate(doc.tables, start=1):
    print(f"TABLE {table_index}: {len(table.rows)} rows x {len(table.columns)} cols")
    for row_index, row in enumerate(table.rows[:8], start=1):
        cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
        print(f"  R{row_index}: {' | '.join(cells)}")
