from pathlib import Path
from docx import Document


paths = [
    Path(r"C:\Users\chiqu\Downloads\Plantilla Documentacion (2).docx"),
    Path(r"C:\Users\chiqu\tecwebSistemaGym\SistemaGym.Api\DocumentacionFinal_SistemaGym_Actualizada.docx"),
]

for path in paths:
    print(f"\n===== {path} =====")
    doc = Document(path)
    in_section_5 = False
    for i, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
        if text.startswith("5.") or text == "5. Implementación" or text == "5. Implementacion":
            in_section_5 = True
        elif in_section_5 and (text.startswith("6.") or text == "6. Anexos"):
            break
        if in_section_5:
            print(f"{i:04d} [{paragraph.style.name}] {text}")

    print("TABLES AROUND SECTION 5")
    for table_index, table in enumerate(doc.tables, start=1):
        header = " | ".join(cell.text.replace("\n", " ").strip() for cell in table.rows[0].cells)
        if any(token.lower() in header.lower() for token in ["tecnolog", "método", "metodo", "criterio", "caso", "componente"]):
            print(f"TABLE {table_index}: {len(table.rows)}x{len(table.columns)} :: {header}")
