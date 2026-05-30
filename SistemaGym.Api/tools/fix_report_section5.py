from pathlib import Path
from docx import Document


PATH = Path(r"C:\Users\chiqu\tecwebSistemaGym\SistemaGym.Api\DocumentacionFinal_SistemaGym_Actualizada.docx")

doc = Document(PATH)

step_texts = {
    "5. Ejecutar los endpoints protegidos de clientes, planes, membresías y pagos.": "1. Abrir Swagger en la URL pública /swagger.",
    "4. Presionar Authorize en Swagger y registrar el valor con el formato Bearer {token}.": "2. Ejecutar POST /api/Token con user=admin@gym.com y password=123456.",
    "3. Copiar el token devuelto por la API.": "3. Copiar el token devuelto por la API.",
    "2. Ejecutar POST /api/Token con user=admin@gym.com y password=123456.": "4. Presionar Authorize en Swagger y registrar el valor con el formato Bearer {token}.",
    "1. Abrir Swagger en la URL pública /swagger.": "5. Ejecutar los endpoints protegidos de clientes, planes, membresías y pagos.",
}

for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    if text in step_texts:
        paragraph.text = step_texts[text]

for index, paragraph in enumerate(doc.paragraphs):
    if paragraph.text.strip() == "(Enlaces y recursos vivos que complementan el documento estático.)":
        previous = doc.paragraphs[index - 1]
        previous.text = "6. Anexos"
        previous.style = doc.styles["Heading 1"]
        break

doc.save(PATH)
print(PATH)
