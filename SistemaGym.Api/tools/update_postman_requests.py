from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement


PATH = Path(r"C:\Users\chiqu\tecwebSistemaGym\SistemaGym.Api\DocumentacionFinal_SistemaGym_Actualizada.docx")
BASE_URL = "https://sistemagym-api-gee7bwdvepgzc0f6.canadacentral-01.azurewebsites.net"

requests = [
    "Autenticación: POST {base}/api/Token",
    "Body: { \"user\": \"admin@gym.com\", \"password\": \"123456\" }",
    "CU-01 Registrar Cliente: POST {base}/api/Cliente/dto/mapper",
    "Body: { \"id\": 0, \"nombre\": \"Carlos\", \"apellido\": \"Rojas\", \"ci\": \"9876543\", \"correo\": \"carlos@gym.com\", \"telefono\": \"77712345\", \"fechaRegistro\": \"2026-05-29T00:00:00\", \"estado\": true }",
    "CU-02 Consultar Clientes: GET {base}/api/Cliente/dto/mapper?pageNumber=1&pageSize=10",
    "CU-02 Consultar Clientes con filtro: GET {base}/api/Cliente/dto/mapper?nombre=Juan&pageNumber=1&pageSize=10",
    "CU-03 Actualizar Cliente: PUT {base}/api/Cliente/dto/mapper/1",
    "Body: { \"id\": 1, \"nombre\": \"Juan\", \"apellido\": \"Perez\", \"ci\": \"1234567\", \"correo\": \"juan.actualizado@gmail.com\", \"telefono\": \"70000001\", \"fechaRegistro\": \"2026-05-29T00:00:00\", \"estado\": true }",
    "CU-04 Eliminar Cliente: DELETE {base}/api/Cliente/dto/1",
    "CU-05 Crear Plan de Membresía: POST {base}/api/PlanMembresia/dto/mapper",
    "Body: { \"id\": 0, \"nombrePlan\": \"Plan Mensual\", \"descripcion\": \"Acceso completo al gimnasio por 30 dias\", \"duracionDias\": 30, \"precio\": 150.00, \"estado\": true }",
    "CU-06 Consultar Planes de Membresía: GET {base}/api/PlanMembresia/dto/mapper?pageNumber=1&pageSize=10",
    "CU-06 Consultar Planes por precio: GET {base}/api/PlanMembresia/dto/mapper?precioMin=100&precioMax=300&pageNumber=1&pageSize=10",
    "CU-06 Consulta Dapper de Planes: GET {base}/api/PlanMembresia/dto/mapper/dapper?limit=10",
    "CU-07 Asignar Membresía a Cliente: POST {base}/api/Membresia/dto/mapper",
    "Body: { \"id\": 0, \"clienteId\": 1, \"planMembresiaId\": 1, \"fechaInicio\": \"2026-05-29T00:00:00\", \"fechaFin\": \"2026-06-28T00:00:00\", \"estado\": true }",
    "CU-08 Consultar Membresías: GET {base}/api/Membresia/dto/mapper?pageNumber=1&pageSize=10",
    "CU-08 Consultar Membresías con filtro: GET {base}/api/Membresia/dto/mapper?clienteId=1&estado=true&pageNumber=1&pageSize=10",
    "CU-08 Consulta Dapper de Membresías: GET {base}/api/Membresia/dto/mapper/dapper?limit=10",
    "CU-09 Registrar Pago de Membresía: POST {base}/api/Pago/dto/mapper",
    "Body: { \"id\": 0, \"membresiaId\": 1, \"monto\": 150.00, \"fechaPago\": \"2026-05-29T00:00:00\", \"metodoPago\": \"Efectivo\", \"estado\": true }",
    "CU-10 Consultar Pagos de Membresía: GET {base}/api/Pago/dto/mapper?pageNumber=1&pageSize=10",
    "CU-10 Consultar Pago por ID: GET {base}/api/Pago/dto/mapper/1",
]
requests = [item.replace("{base}", BASE_URL) for item in requests]


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p = new_p
    new_paragraph._element = new_p
    new_paragraph.text = text
    new_paragraph.paragraph_format.space_after = Pt(3)
    return new_paragraph


doc = Document(PATH)

heading = None
heading_64 = None
for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    if text == "6.3 Colección de Requests en Postman":
        heading = paragraph
    elif text.startswith("6.4"):
        heading_64 = paragraph
        break

if heading is None or heading_64 is None:
    raise RuntimeError("No se encontró el bloque 6.3/6.4 para actualizar.")

for paragraph in list(doc.paragraphs):
    text = paragraph.text.strip()
    if (
        text.startswith("Autenticación:")
        or text.startswith("CU-")
        or text.startswith("Body:")
        or text.startswith("Listado de requests realizados mediante Postman/Swagger")
    ):
        delete_paragraph(paragraph)

anchor = heading
anchor = insert_paragraph_after(
    anchor,
    "Listado de requests realizados mediante Postman/Swagger. No se adjunta ni comparte colección; se documentan las rutas ejecutadas, método HTTP y ejemplos de body cuando corresponde.",
)
for text in requests:
    anchor = insert_paragraph_after(anchor, text)

doc.save(PATH)
print(PATH)
